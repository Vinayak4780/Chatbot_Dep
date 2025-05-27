from docx import Document
import re

def preprocess_text(text):
    # Normalize spaces and remove extra blank lines
    text = text.strip()
    text = re.sub(r'\s+', ' ', text)  # Replace multiple spaces/newlines with single space
    return text

def extract_docx_to_txt(input_file, output_file):
    doc = Document(input_file)
    with open(output_file, 'w', encoding='utf-8') as f:
        for para in doc.paragraphs:
            if para.text.strip():
                # If the paragraph has bold text (title)
                if any(run.bold for run in para.runs):
                    f.write("\n=== TITLE: " + para.text.strip() + " ===\n")
                else:
                    f.write(preprocess_text(para.text) + "\n")

        # Handle tables
        for table_index, table in enumerate(doc.tables):
            f.write(f"\n=== TABLE {table_index + 1} ===\n")
            for row in table.rows:
                row_text = '\t'.join(cell.text.strip() for cell in row.cells)
                f.write(preprocess_text(row_text) + "\n")

    print(f"✔️ Successfully saved to {output_file}")

# Usage
input_docx = "/home/ubuntu/DSEU-Website-Main/chatbot/DSEU_Admission Brochure 2025 final(1).docx"
output_txt = "output.txt"
extract_docx_to_txt(input_docx, output_txt)






from langchain.schema import Document
import re

def load_structured_txt(file_path):
    documents = []
    current_title = ""
    current_type = "text"
    buffer = []

    with open(file_path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()

            if not line:
                continue

            # Title detection
            title_match = re.match(r"=== TITLE: (.+?) ===", line)
            if title_match:
                if buffer:
                    documents.append(Document(page_content="\n".join(buffer), metadata={"title": current_title, "type": current_type}))
                    buffer = []
                current_title = title_match.group(1)
                current_type = "text"
                continue

            # Table detection
            table_match = re.match(r"=== TABLE (\d+) ===", line)
            if table_match:
                if buffer:
                    documents.append(Document(page_content="\n".join(buffer), metadata={"title": current_title, "type": current_type}))
                    buffer = []
                current_title = f"Table {table_match.group(1)}"
                current_type = "table"
                continue

            buffer.append(line)

        if buffer:
            documents.append(Document(page_content="\n".join(buffer), metadata={"title": current_title, "type": current_type}))

    return documents
from langchain.document_loaders import TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
import os

# 1. Load your raw text into Documents
loader = TextLoader("output.txt", encoding="utf-8")
raw_docs = loader.load()  # returns List[Document]

# 2. Split long docs into chunks for better retrieval quality
splitter = RecursiveCharacterTextSplitter(
    chunk_size=500,        # max characters per chunk
    chunk_overlap=20 ,      # characters overlap between chunks
    separators=["\n\n", "\n", " ", ""],
)
docs = splitter.split_documents(raw_docs)

# 3. Choose and instantiate your embedding model
embeddings = HuggingFaceEmbeddings(
    model_name="sentence-transformers/all-MiniLM-L6-v2",
    # cache_folder=os.getenv("HF_HOME", "~/.cache/huggingface")  # optional caching
)

# 4. Build the FAISS index
vectorstore = FAISS.from_documents(docs, embeddings)

# 5. Persist the index to disk
output_dir = "data/defaultllm_docs"
os.makedirs(output_dir, exist_ok=True)
vectorstore.save_local(output_dir)

print(f"Indexed {len(docs)} document chunks and saved to '{output_dir}'.")














# from docx import Document as DocxDocument
# import re
# import os
# from langchain.schema import Document as LcDocument
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from langchain.embeddings import HuggingFaceEmbeddings
# from langchain.vectorstores import FAISS


# def preprocess_text(text):
#     """
#     Normalize spaces and remove extra blank lines in a text string.
#     """
#     text = text.strip()
#     text = re.sub(r"\s+", ' ', text)
#     return text


# def extract_docx_to_txt(input_file: str, output_file: str):
#     """
#     Extract paragraphs and tables from a DOCX file into a structured TXT file.
#     Paragraphs with bold runs are marked as titles.
#     Tables are written row-by-row with tab-separated columns.
#     """
#     doc = DocxDocument(input_file)
#     with open(output_file, 'w', encoding='utf-8') as f:
#         # Process paragraphs
#         for para in doc.paragraphs:
#             if para.text.strip():
#                 if any(run.bold for run in para.runs):
#                     f.write(f"\n=== TITLE: {para.text.strip()} ===\n")
#                 else:
#                     f.write(preprocess_text(para.text) + "\n")
#         # Process tables
#         for idx, table in enumerate(doc.tables, start=1):
#             f.write(f"\n=== TABLE {idx} ===\n")
#             for row in table.rows:
#                 row_text = '\t'.join(cell.text.strip() for cell in row.cells)
#                 f.write(preprocess_text(row_text) + "\n")

#     print(f"✔️ Successfully saved to {output_file}")


# def load_structured_txt(file_path: str) -> list:
#     """
#     Read the structured TXT file, returning a list of LangChain Documents
#     with metadata {'title': ..., 'type': 'text'|'table'}.
#     """
#     docs = []
#     current_title = ''
#     current_type = 'text'
#     buffer = []

#     with open(file_path, 'r', encoding='utf-8') as f:
#         for line in f:
#             line = line.strip()
#             if not line:
#                 continue
#             title_match = re.match(r"=== TITLE: (.+?) ===", line)
#             table_match = re.match(r"=== TABLE (\d+) ===", line)

#             if title_match:
#                 if buffer:
#                     docs.append(LcDocument(page_content="\n".join(buffer), metadata={
#                         'title': current_title,
#                         'type': current_type
#                     }))
#                     buffer = []
#                 current_title = title_match.group(1)
#                 current_type = 'text'
#                 continue

#             if table_match:
#                 if buffer:
#                     docs.append(LcDocument(page_content="\n".join(buffer), metadata={
#                         'title': current_title,
#                         'type': current_type
#                     }))
#                     buffer = []
#                 current_title = f"Table {table_match.group(1)}"
#                 current_type = 'table'
#                 continue

#             buffer.append(line)

#         # Append last buffered segment
#         if buffer:
#             docs.append(LcDocument(page_content="\n".join(buffer), metadata={
#                 'title': current_title,
#                 'type': current_type
#             }))
#     return docs


# if __name__ == '__main__':
#     # Paths
#     input_docx = "/home/ubuntu/DSEU-Website-Main/chatbot/DSEU_Admission Brochure 2025 final(1).docx"
#     structured_txt = "output.txt"
#     output_dir = "data/defaultllm_docs"

#     # Step 1: Extract from DOCX
#     extract_docx_to_txt(input_docx, structured_txt)

#     # Step 2: Load structured docs
#     raw_docs = load_structured_txt(structured_txt)

#     # Prepare splitter for text
#     splitter = RecursiveCharacterTextSplitter(
#         chunk_size=1000,
#         chunk_overlap=200,
#         separators=["\n\n", "\n", " ", ""]
#     )

#     # Step 3: Process docs into chunks
#     processed_chunks = []
#     for doc in raw_docs:
#         if doc.metadata.get('type') == 'text':
#             # Split long text into chunks
#             processed_chunks.extend(splitter.split_documents([doc]))
#         else:
#             # For tables, split into individual rows
#             rows = doc.page_content.split("\n")
#             for row in rows:
#                 if row.strip():
#                     processed_chunks.append(
#                         LcDocument(page_content=row,
#                                    metadata={
#                                        'title': doc.metadata.get('title'),
#                                        'type': 'table_row'
#                                    })
#                     )

#     # Step 4: Build embeddings and FAISS index
#     embeddings = HuggingFaceEmbeddings(
#         model_name="sentence-transformers/all-MiniLM-L6-v2"
#     )
#     os.makedirs(output_dir, exist_ok=True)
#     vectorstore = FAISS.from_documents(processed_chunks, embeddings)
#     vectorstore.save_local(output_dir)

#     print(f"Indexed {len(processed_chunks)} chunks (text & table rows) to '{output_dir}'")








